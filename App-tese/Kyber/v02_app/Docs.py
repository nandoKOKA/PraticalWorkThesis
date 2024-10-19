import os
import random
from docx import Document
from openpyxl import Workbook
from openpyxl import load_workbook
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import random
import string

def generate_random_word_file(file_path, num_paragraphs):
    document = Document()
    for _ in range(num_paragraphs):
        paragraph = document.add_paragraph()
        num_sentences = random.randint(1, 5)
        for _ in range(num_sentences):
            sentence = ''.join(random.choices('abcdefghijklmnopqrstuvwxyz ', k=random.randint(5, 15)))
            paragraph.add_run(sentence)
        paragraph.add_run('\n')
    document.save(file_path)

def generate_random_excel_file(file_path, num_rows, num_columns):
    workbook = Workbook()
    sheet = workbook.active
    for _ in range(num_rows):
        row_data = [random.randint(1, 100) for _ in range(num_columns)]
        sheet.append(row_data)
    workbook.save(file_path)

def generate_random_text(length):
    return ''.join(random.choices(string.ascii_letters + string.digits, k=length))

def generate_random_pdf(file_path, num_paragraphs):
    doc = SimpleDocTemplate(file_path, pagesize=letter)
    styles = getSampleStyleSheet()
    paragraphs = []

    for _ in range(num_paragraphs):
        random_text = generate_random_text(random.randint(50, 200))
        paragraph = Paragraph(random_text, styles["Normal"])
        paragraphs.append(paragraph)

    doc.build(paragraphs)

# Exemplo de uso
generate_random_pdf("random_document.pdf", 1000000)
generate_random_word_file("random_word_file.docx", 5)
generate_random_excel_file("random_excel_file.xlsx", 5, 5)

